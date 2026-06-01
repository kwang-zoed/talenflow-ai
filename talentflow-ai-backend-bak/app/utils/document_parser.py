import io

def extract_text_from_bytes(file_content: bytes, filename: str) -> str:
    import pdfplumber
    full_text = ""
    if filename.lower().endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text += text + '\n'
    elif filename.lower().endswith(".docx"):
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(file_content))
        parts = [p.text.strip() for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text.strip())
        full_text = '\n'.join(parts)
    elif filename.lower().endswith(".txt"):
        if isinstance(file_content, bytes):
            full_text = file_content.decode('utf-8', errors='ignore')
        else:
            full_text = str(file_content)
    return full_text

def extract_text_from_file(contents, filename):
    return extract_text_from_bytes(contents, filename)
